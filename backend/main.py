import io
import os
from typing import List, Optional

import google.generativeai as genai
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from PIL import Image
import httpx
from pydantic import BaseModel
from supabase import Client, create_client
from docx import Document


load_dotenv()


SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
SUPABASE_BUCKET_IMAGES = os.getenv("SUPABASE_BUCKET_IMAGES", "defect-images")
SUPABASE_BUCKET_REPORTS = os.getenv("SUPABASE_BUCKET_REPORTS", "defect-reports")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
EXAMINER_INSTRUCTIONS = os.getenv(
    "EXAMINER_INSTRUCTIONS",
    "You are a construction examiner following GOST standards. "
    "Given a structural photo, describe visible defects and assign a status category "
    "(A - no defect, B - minor, C - serious, D - critical).",
)


def _require_supabase() -> Client:
    if not (SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY):
        raise HTTPException(status_code=500, detail="Supabase is not configured.")
    return create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)  # type: ignore[arg-type]


def _require_gemini_model() -> genai.GenerativeModel:
    if not GEMINI_API_KEY:
        raise HTTPException(status_code=500, detail="Gemini API key is not configured.")
    genai.configure(api_key=GEMINI_API_KEY)
    return genai.GenerativeModel("gemini-1.5-flash")


class SurveyCreate(BaseModel):
    user_id: str
    name: str
    industry_gost: str


class Survey(BaseModel):
    id: str
    user_id: str
    name: str
    industry_gost: str


class DefectAnalyzeRequest(BaseModel):
    survey_id: str
    image_url: str
    axis: str  # "X" or "Y"
    construction_type: str  # "Concrete" | "Brick" | "Metal" | "Roof"
    location: Optional[str] = None


class Defect(BaseModel):
    id: str
    survey_id: str
    image_url: str
    axis: str
    construction_type: str
    location: Optional[str]
    description: str
    status_category: str


app = FastAPI(title="Industrial Defect Examiner API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/surveys", response_model=Survey)
async def create_survey(payload: SurveyCreate):
    supabase = _require_supabase()
    data = {
        "user_id": payload.user_id,
        "name": payload.name,
        "industry_gost": payload.industry_gost,
    }
    try:
        res = supabase.table("surveys").insert(data).execute()
    except Exception as e:  # pragma: no cover - supabase client errors
        raise HTTPException(status_code=500, detail=f"Failed to create survey: {e}")

    if not res.data:
        raise HTTPException(status_code=500, detail="Supabase did not return survey.")

    row = res.data[0]
    return Survey(
        id=str(row["id"]),
        user_id=row["user_id"],
        name=row["name"],
        industry_gost=row["industry_gost"],
    )


@app.get("/surveys", response_model=List[Survey])
async def list_surveys(user_id: str):
    supabase = _require_supabase()
    try:
        res = supabase.table("surveys").select("*").eq("user_id", user_id).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch surveys: {e}")

    return [
        Survey(
            id=str(row["id"]),
            user_id=row["user_id"],
            name=row["name"],
            industry_gost=row["industry_gost"],
        )
        for row in (res.data or [])
    ]


async def _fetch_image_as_pil(url: str) -> Image.Image:
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(url)
        if resp.status_code != 200:
            raise HTTPException(status_code=400, detail=f"Could not fetch image: {resp.status_code}")
        return Image.open(io.BytesIO(resp.content)).convert("RGB")


def _build_system_prompt() -> str:
    return (
        f"{EXAMINER_INSTRUCTIONS}\n\n"
        "Return a concise JSON object with keys:\n"
        '"description": string (defect description using GOST terminology when possible),\n'
        '"status_category": string (one of: A, B, C, D).\n'
        "Respond ONLY with valid JSON."
    )


async def _analyze_defect_with_gemini(
    image: Image.Image,
    axis: str,
    construction_type: str,
    location: Optional[str],
) -> dict:
    user_prompt = (
        "Analyze this construction element photo.\n"
        f"Axis: {axis}\n"
        f"Construction type: {construction_type}\n"
        f"Location: {location or 'unspecified'}\n"
    )
    gemini_model = _require_gemini_model()
    result = gemini_model.generate_content(
        [
            _build_system_prompt(),
            user_prompt,
            image,
        ],
        generation_config={"response_mime_type": "application/json"},
    )

    try:
        # `result.text` should contain JSON string
        import json

        data = json.loads(result.text or "{}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Gemini response: {e}")

    description = data.get("description") or "No description provided."
    status_category = data.get("status_category") or "B"

    return {"description": description, "status_category": status_category}


@app.post("/defects/analyze", response_model=Defect)
async def analyze_defect(payload: DefectAnalyzeRequest):
    supabase = _require_supabase()
    image = await _fetch_image_as_pil(payload.image_url)
    gemini_result = await _analyze_defect_with_gemini(
        image=image,
        axis=payload.axis,
        construction_type=payload.construction_type,
        location=payload.location,
    )

    defect_row = {
        "survey_id": payload.survey_id,
        "image_url": payload.image_url,
        "axis": payload.axis,
        "construction_type": payload.construction_type,
        "location": payload.location,
        "description": gemini_result["description"],
        "status_category": gemini_result["status_category"],
    }

    try:
        res = supabase.table("defects").insert(defect_row).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save defect: {e}")

    if not res.data:
        raise HTTPException(status_code=500, detail="Supabase did not return defect.")

    row = res.data[0]
    return Defect(
        id=str(row["id"]),
        survey_id=row["survey_id"],
        image_url=row["image_url"],
        axis=row["axis"],
        construction_type=row["construction_type"],
        location=row.get("location"),
        description=row["description"],
        status_category=row["status_category"],
    )


@app.get("/defects/{survey_id}", response_model=List[Defect])
async def list_defects_for_survey(survey_id: str):
    supabase = _require_supabase()
    try:
        res = supabase.table("defects").select("*").eq("survey_id", survey_id).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch defects: {e}")

    defects = res.data or []
    return [
        Defect(
            id=str(row["id"]),
            survey_id=row["survey_id"],
            image_url=row["image_url"],
            axis=row["axis"],
            construction_type=row["construction_type"],
            location=row.get("location"),
            description=row["description"],
            status_category=row["status_category"],
        )
        for row in defects
    ]


@app.get("/reports/{survey_id}")
async def generate_report(survey_id: str):
    supabase = _require_supabase()
    try:
        defects_res = supabase.table("defects").select("*").eq("survey_id", survey_id).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch defects: {e}")

    defects = defects_res.data or []

    if not defects:
        raise HTTPException(status_code=404, detail="No defects found for survey.")

    # Build DOCX in memory
    doc = Document()
    doc.add_heading("Construction Defect Report", level=1)
    doc.add_paragraph(f"Survey ID: {survey_id}")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Photo"
    hdr_cells[1].text = "Location"
    hdr_cells[2].text = "Description of the defect"
    hdr_cells[3].text = "Status category"

    for d in defects:
        row_cells = table.add_row().cells
        row_cells[0].text = d.get("image_url", "")
        row_cells[1].text = d.get("location") or ""
        row_cells[2].text = d.get("description") or ""
        row_cells[3].text = d.get("status_category") or ""

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"survey-{survey_id}-report.docx"

    # Optionally upload to Supabase storage
    try:
        supabase.storage.from_(SUPABASE_BUCKET_REPORTS).upload(
            file=buffer.getvalue(),
            path=filename,
            file_options={"content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
        )
    except Exception as e:
        # Non-fatal: report is still streamed back to client
        print(f"WARNING: Failed to upload report to Supabase: {e}")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# To run locally:
# uvicorn main:app --reload --host 0.0.0.0 --port 8000

